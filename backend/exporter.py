import json

from docx import Document
from xml.sax.saxutils import escape
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

from qa_utils import summarize_qa, qa_flag_text, is_warning_line
import risk_map

# ---------------------------------------------------------------------------
# Shared helpers for the latest features: risk priority + framework mapping
# (risk_map) and the retest workflow (retest_* fields on each finding).
# All are deterministic/offline and never raise on malformed input.
# ---------------------------------------------------------------------------
_RETEST_OPEN_STATES = ("Open", "Partially Fixed", "Regressed")
_SEVERITY_ORDER = ["Critical", "High", "Medium", "Low", "Informational"]
_PRIORITY_ORDER = ["Urgent", "High", "Moderate", "Low"]

# Risk-priority colours (DOCX RGB and PDF hex).
_RISK_RGB = {
    "Urgent": (0xC0, 0x28, 0x28),
    "High": (0xB9, 0x5C, 0x00),
    "Moderate": (0x2B, 0x6C, 0xB0),
    "Low": (0x4A, 0x55, 0x68),
}
_RISK_HEX = {
    "Urgent": "#c02828",
    "High": "#b95c00",
    "Moderate": "#2b6cb0",
    "Low": "#4a5568",
}
# XLSX priority cell fills.
_RISK_XLSX_FILL = {
    "Urgent": PatternFill(start_color="F8D7DA", end_color="F8D7DA", fill_type="solid"),
    "High": PatternFill(start_color="FCE5CD", end_color="FCE5CD", fill_type="solid"),
    "Moderate": PatternFill(start_color="DDE7F3", end_color="DDE7F3", fill_type="solid"),
    "Low": PatternFill(start_color="EDF2F7", end_color="EDF2F7", fill_type="solid"),
}


def _retested(finding):
    """True if the finding has at least one recorded retest outcome."""
    status = str((finding or {}).get("retest_status") or "").strip()
    try:
        rounds = int((finding or {}).get("retest_round") or 0)
    except (TypeError, ValueError):
        rounds = 0
    return rounds > 0 or (bool(status) and status != "Not Retested")


def _retest_history(finding):
    """Parse the retest_history JSON on a finding into a list (never raises)."""
    raw = (finding or {}).get("retest_history")
    if not raw:
        return []
    try:
        data = json.loads(raw) if isinstance(raw, str) else raw
        return data if isinstance(data, list) else []
    except Exception:
        return []


def _retest_summary(findings):
    """Remediation rollup mirrored from DatabaseManager.get_retest_summary so the
    exporters stay self-contained (no DB handle needed)."""
    total = len(findings)
    counts = {}
    residual = {}
    for f in findings:
        state = f.get("retest_status") or "Not Retested"
        counts[state] = counts.get(state, 0) + 1
        if state in _RETEST_OPEN_STATES:
            sev = f.get("original_severity") or f.get("severity") or "Unknown"
            residual[sev] = residual.get(sev, 0) + 1
    fixed = counts.get("Fixed", 0)
    accepted = counts.get("Accepted Risk", 0)
    not_retested = counts.get("Not Retested", 0)
    open_like = sum(counts.get(s, 0) for s in _RETEST_OPEN_STATES)
    retested = total - not_retested
    rate = (fixed / retested * 100.0) if retested else 0.0
    return {
        "total": total, "retested": retested, "not_retested": not_retested,
        "fixed": fixed, "accepted": accepted, "open": open_like,
        "counts": counts, "remediation_rate": rate,
        "residual_by_severity": residual, "any_retested": retested > 0,
    }


def _risk_summary(findings):
    """Counts of findings by risk priority and by OWASP 2025 category."""
    priorities = {}
    owasp = {}
    unmapped = 0
    for f in findings:
        pr = risk_map.compute_risk_priority(f)["priority"]
        priorities[pr] = priorities.get(pr, 0) + 1
        fw = risk_map.map_frameworks(f)
        if fw.get("mapped") and fw.get("owasp"):
            owasp[fw["owasp"]] = owasp.get(fw["owasp"], 0) + 1
        else:
            unmapped += 1
    return {"priorities": priorities, "owasp": owasp, "unmapped": unmapped}


def _order_severity(mapping):
    return [s for s in _SEVERITY_ORDER if s in mapping] + \
           [s for s in mapping if s not in _SEVERITY_ORDER]


# ===========================================================================
# DOCX
# ===========================================================================
def _docx_table(doc, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    return table


def export_to_docx(project: dict, findings: list, exec_summary: str, methodology: str, filepath: str):
    doc = Document()

    # Title Section
    title = doc.add_paragraph()
    title.add_run(f"Penetration Testing Assessment Report\nClient: {project.get('client', '')}")
    title.style.font.name = "Arial"
    title.style.font.size = Pt(24)
    title.style.font.bold = True

    # Project Metadata
    doc.add_heading("Project Overview", level=1)
    p = doc.add_paragraph()
    p.add_run(f"Project Name: {project.get('name', '')}\n")
    p.add_run(f"Assessor: {project.get('tester', '')}\n")
    p.add_run(f"Scope Window: {project.get('start_date', '')} to {project.get('end_date', '')}\n")
    p.add_run(f"Target Scope: {project.get('scope', '')}\n")

    # Executive Summary & Methodology
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(exec_summary if exec_summary else "No summary provided.")

    doc.add_heading("Assessment Methodology", level=1)
    doc.add_paragraph(methodology if methodology else "Standard OWASP and PTES methodology was followed.")

    # Findings Table Summary (now includes risk priority + retest status)
    doc.add_heading("Summary of Findings", level=1)
    table = _docx_table(doc, ["Finding Title", "Severity", "Risk Priority", "Status", "Retest"])
    for item in findings:
        priority = risk_map.compute_risk_priority(item)["priority"]
        cells = table.add_row().cells
        cells[0].text = str(item.get("title") or "")
        cells[1].text = str(item.get("severity") or "")
        cells[2].text = priority
        cells[3].text = str(item.get("status") or "Draft")
        cells[4].text = str(item.get("retest_status") or "Not Retested")

    # Risk Prioritization Summary
    risk_sum = _risk_summary(findings)
    doc.add_heading("Risk Prioritization Summary", level=1)
    doc.add_paragraph(
        "Risk-based prioritization blends CVSS with EPSS exploit probability, CISA KEV "
        "active-exploitation status, and the affected environment. It is distinct from raw CVSS severity."
    )
    prio_table = _docx_table(doc, ["Risk Priority", "Findings"])
    for pr in _PRIORITY_ORDER:
        if risk_sum["priorities"].get(pr):
            cells = prio_table.add_row().cells
            cells[0].text = pr
            cells[1].text = str(risk_sum["priorities"][pr])

    if risk_sum["owasp"] or risk_sum["unmapped"]:
        doc.add_heading("OWASP Top 10:2025 Coverage", level=2)
        owasp_table = _docx_table(doc, ["Category", "Findings"])
        for cat in sorted(risk_sum["owasp"]):
            cells = owasp_table.add_row().cells
            cells[0].text = cat
            cells[1].text = str(risk_sum["owasp"][cat])
        if risk_sum["unmapped"]:
            cells = owasp_table.add_row().cells
            cells[0].text = "Unmapped (assign manually)"
            cells[1].text = str(risk_sum["unmapped"])

    # Remediation Status Summary (only when a retest has been recorded)
    retest_sum = _retest_summary(findings)
    if retest_sum["any_retested"]:
        doc.add_heading("Remediation Status Summary", level=1)
        doc.add_paragraph(
            f"Findings: {retest_sum['total']} | Retested: {retest_sum['retested']} | "
            f"Fixed: {retest_sum['fixed']} | Still open: {retest_sum['open']} | "
            f"Accepted risk: {retest_sum['accepted']} | Not retested: {retest_sum['not_retested']} | "
            f"Remediation rate: {retest_sum['remediation_rate']:.0f}%"
        )
        if retest_sum["residual_by_severity"]:
            doc.add_heading("Residual Risk (Open Findings by Original Severity)", level=2)
            res_table = _docx_table(doc, ["Severity", "Open"])
            for sev in _order_severity(retest_sum["residual_by_severity"]):
                cells = res_table.add_row().cells
                cells[0].text = str(sev)
                cells[1].text = str(retest_sum["residual_by_severity"][sev])

    # Detailed Findings Breakouts
    doc.add_heading("Detailed Vulnerability Analysis", level=1)
    for f in findings:
        doc.add_heading(f"{f.get('title')} - ({f.get('severity')})", level=2)
        doc.add_paragraph(f"Category: {f.get('category', 'N/A')} | Status: {f.get('status', 'N/A')} | Environment: {f.get('environment', 'N/A')}")
        doc.add_paragraph(f"Affected Host: {f.get('affected_host', 'N/A')} | Affected URL: {f.get('affected_url', 'N/A')}")
        doc.add_paragraph(f"HTTP Method: {f.get('http_method', 'N/A')} | Parameter: {f.get('parameter', 'N/A')}")
        doc.add_paragraph(f"CWE: {f.get('cwe')}  |  CVSS Matrix: {f.get('cvss')}")

        flags = qa_flag_text(f)
        if flags:
            banner = doc.add_paragraph()
            run = banner.add_run(f"VERIFICATION FLAGS: {flags}")
            run.bold = True
            run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)

        # Risk priority + framework mapping
        assessment = risk_map.assess(f)
        risk = assessment["risk"]
        fw = assessment["frameworks"]
        doc.add_heading("Risk Priority & Framework Mapping", level=3)
        rp = doc.add_paragraph()
        run = rp.add_run(f"Risk priority: {risk['priority']} ({risk['score']}/100)")
        run.bold = True
        rgb = _RISK_RGB.get(risk["priority"])
        if rgb:
            run.font.color.rgb = RGBColor(*rgb)
        if risk["rationale"]:
            doc.add_paragraph(f"Rationale: {risk['rationale']}")
        if fw.get("mapped"):
            doc.add_paragraph(f"OWASP Top 10:2025: {fw['owasp']}")
            doc.add_paragraph(f"PCI DSS: {fw['pci']}")
            doc.add_paragraph(f"NIST SP 800-53: {fw['nist']}")
            if fw.get("attack"):
                doc.add_paragraph(f"MITRE ATT&CK: {fw['attack']}")
            note = doc.add_paragraph("Framework mappings are indicative; confirm against the engagement scope.")
            note.runs[0].italic = True
        else:
            doc.add_paragraph(f"Framework mapping: unmapped ({fw.get('basis', '')}) - assign manually.")

        doc.add_heading("Description", level=3)
        doc.add_paragraph(f.get("description"))

        doc.add_heading("Evidence Log / Artifacts", level=3)
        doc.add_paragraph(f.get("evidence"))
        if f.get("evidence_files"):
            doc.add_paragraph(f"Evidence Files: {f.get('evidence_files')}")

        doc.add_heading("Business & Technical Impact", level=3)
        doc.add_paragraph(f.get("impact"))

        doc.add_heading("Exploitation Scenario", level=3)
        doc.add_paragraph(f.get("scenario"))

        doc.add_heading("Reproduction Vectors", level=3)
        doc.add_paragraph(f.get("steps"))

        doc.add_heading("Remediation & Hardening Roadmap", level=3)
        doc.add_paragraph(f.get("remediation"))

        # Retest detail (only when retested)
        if _retested(f):
            doc.add_heading("Retest", level=3)
            doc.add_paragraph(
                f"Retest status: {f.get('retest_status') or 'N/A'} | Round: {f.get('retest_round') or 0} | "
                f"Date: {f.get('retest_date') or 'N/A'} | Retester: {f.get('retester') or 'N/A'}"
            )
            doc.add_paragraph(
                f"Original severity: {f.get('original_severity') or f.get('severity') or 'N/A'} | "
                f"Current severity: {f.get('severity') or 'N/A'} | "
                f"First found: {f.get('first_found_date') or f.get('created_at') or 'N/A'}"
            )
            if f.get("retest_evidence"):
                doc.add_paragraph(f"Retest evidence: {f.get('retest_evidence')}")
            history = _retest_history(f)
            if history:
                doc.add_paragraph("Retest history:")
                for ev in history:
                    line = f"  R{ev.get('round')} - {ev.get('date')} - {ev.get('status')}"
                    if ev.get("retester"):
                        line += f" - {ev.get('retester')}"
                    if ev.get("note"):
                        line += f" - {ev.get('note')}"
                    doc.add_paragraph(line)

        doc.add_heading("False Positive Check / Retest Notes", level=3)
        doc.add_paragraph(f.get("fp_checks") or "")
        if f.get("retest_notes"):
            doc.add_paragraph(f"Retest Notes: {f.get('retest_notes')}")

        qa_text = (f.get("additional_remarks") or "").strip()
        if qa_text:
            doc.add_heading("Automated QA / Verification", level=3)
            for line in qa_text.splitlines():
                if not line.strip():
                    continue
                para = doc.add_paragraph()
                run = para.add_run(line)
                if is_warning_line(line):
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)

        doc.add_heading("References", level=3)
        doc.add_paragraph(f.get("references_data") or f.get("references", ""))

    doc.save(filepath)


# ===========================================================================
# PDF
# ===========================================================================
def _pdf_text(value) -> str:
    """Escape raw Burp evidence before sending it to ReportLab Paragraph.
    ReportLab parses <tag> as markup, so raw HTML responses must be escaped."""
    if value is None:
        return ""
    text = str(value).replace("\r\n", "\n").replace("\r", "\n")
    text = escape(text, entities={"'": "&#39;", '"': "&quot;"})
    return text.replace("\n", "<br/>")


def _para(value, style):
    return Paragraph(_pdf_text(value), style)


def export_to_pdf(project: dict, findings: list, exec_summary: str, methodology: str, filepath: str):
    doc = SimpleDocTemplate(
        filepath, pagesize=letter,
        rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36,
    )
    styles = getSampleStyleSheet()
    story = []

    title_style = ParagraphStyle("DocTitle", parent=styles["Heading1"], fontSize=22,
                                 textColor=colors.HexColor("#1a365d"), spaceAfter=15, wordWrap="CJK")
    h1_style = ParagraphStyle("H1", parent=styles["Heading2"], fontSize=16,
                              textColor=colors.HexColor("#2b6cb0"), spaceBefore=12, spaceAfter=6, wordWrap="CJK")
    h2_style = ParagraphStyle("H2", parent=styles["Heading3"], fontSize=12,
                              textColor=colors.HexColor("#2d3748"), spaceBefore=8, spaceAfter=4, wordWrap="CJK")
    body_style = ParagraphStyle("BodySafe", parent=styles["Normal"], fontSize=9,
                                spaceAfter=6, leading=12, wordWrap="CJK")
    evidence_style = ParagraphStyle("EvidenceSafe", parent=styles["Code"], fontSize=7,
                                    leading=9, spaceAfter=8, wordWrap="CJK")
    qa_flag_style = ParagraphStyle("QAFlag", parent=styles["Normal"], fontSize=10,
                                   textColor=colors.HexColor("#b00000"), spaceBefore=4, spaceAfter=6, wordWrap="CJK")
    qa_warn_style = ParagraphStyle("QAWarn", parent=styles["Normal"], fontSize=9,
                                   textColor=colors.HexColor("#b00000"), leading=12, spaceAfter=4, wordWrap="CJK")

    summary_table_style = TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#edf2f7")),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("PADDING", (0, 0), (-1, -1), 5),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ])

    def priority_para(priority):
        hexc = _RISK_HEX.get(priority, "#2d3748")
        return Paragraph(f"<font color='{hexc}'><b>{priority}</b></font>", body_style)

    story.append(_para(f"Security Assessment Report: {project.get('name', '')}", title_style))
    story.append(Paragraph(
        f"<b>Client:</b> {_pdf_text(project.get('client', ''))} | "
        f"<b>Tester:</b> {_pdf_text(project.get('tester', ''))}", body_style))
    story.append(Paragraph(
        f"<b>Timeline:</b> {_pdf_text(project.get('start_date', ''))} - "
        f"{_pdf_text(project.get('end_date', ''))}", body_style))
    story.append(Spacer(1, 12))

    story.append(Paragraph("Executive Summary", h1_style))
    story.append(_para(exec_summary or "N/A", body_style))

    story.append(Paragraph("Methodology", h1_style))
    story.append(_para(methodology or "N/A", body_style))

    story.append(Spacer(1, 12))
    story.append(Paragraph("Identified Vulnerabilities", h1_style))

    table_data = [[
        Paragraph("<b>Finding Title</b>", body_style),
        Paragraph("<b>Severity</b>", body_style),
        Paragraph("<b>Risk</b>", body_style),
        Paragraph("<b>Status</b>", body_style),
        Paragraph("<b>CWE</b>", body_style),
    ]]
    for f in findings:
        priority = risk_map.compute_risk_priority(f)["priority"]
        table_data.append([
            _para(f.get("title", ""), body_style),
            _para(f.get("severity", ""), body_style),
            priority_para(priority),
            _para(f.get("status", ""), body_style),
            _para(f.get("cwe", ""), body_style),
        ])
    t = Table(table_data, colWidths=[175, 60, 70, 70, 125], repeatRows=1)
    t.setStyle(summary_table_style)
    story.append(t)
    story.append(Spacer(1, 15))

    # Risk Prioritization Summary
    risk_sum = _risk_summary(findings)
    story.append(Paragraph("Risk Prioritization Summary", h1_style))
    story.append(_para(
        "Risk-based prioritization blends CVSS with EPSS exploit probability, CISA KEV "
        "active-exploitation status, and the affected environment. It is distinct from raw CVSS severity.",
        body_style))
    prio_rows = [[Paragraph("<b>Risk Priority</b>", body_style), Paragraph("<b>Findings</b>", body_style)]]
    for pr in _PRIORITY_ORDER:
        if risk_sum["priorities"].get(pr):
            prio_rows.append([priority_para(pr), _para(str(risk_sum["priorities"][pr]), body_style)])
    prio_table = Table(prio_rows, colWidths=[150, 80], repeatRows=1)
    prio_table.setStyle(summary_table_style)
    story.append(prio_table)
    story.append(Spacer(1, 10))

    if risk_sum["owasp"] or risk_sum["unmapped"]:
        story.append(Paragraph("OWASP Top 10:2025 Coverage", h2_style))
        owasp_rows = [[Paragraph("<b>Category</b>", body_style), Paragraph("<b>Findings</b>", body_style)]]
        for cat in sorted(risk_sum["owasp"]):
            owasp_rows.append([_para(cat, body_style), _para(str(risk_sum["owasp"][cat]), body_style)])
        if risk_sum["unmapped"]:
            owasp_rows.append([_para("Unmapped (assign manually)", body_style), _para(str(risk_sum["unmapped"]), body_style)])
        owasp_table = Table(owasp_rows, colWidths=[320, 80], repeatRows=1)
        owasp_table.setStyle(summary_table_style)
        story.append(owasp_table)
        story.append(Spacer(1, 10))

    # Remediation Status Summary (only when a retest has been recorded)
    retest_sum = _retest_summary(findings)
    if retest_sum["any_retested"]:
        story.append(Paragraph("Remediation Status Summary", h1_style))
        story.append(_para(
            f"Findings: {retest_sum['total']} | Retested: {retest_sum['retested']} | "
            f"Fixed: {retest_sum['fixed']} | Still open: {retest_sum['open']} | "
            f"Accepted risk: {retest_sum['accepted']} | Not retested: {retest_sum['not_retested']} | "
            f"Remediation rate: {retest_sum['remediation_rate']:.0f}%", body_style))
        if retest_sum["residual_by_severity"]:
            story.append(Paragraph("Residual Risk (Open Findings by Original Severity)", h2_style))
            res_rows = [[Paragraph("<b>Severity</b>", body_style), Paragraph("<b>Open</b>", body_style)]]
            for sev in _order_severity(retest_sum["residual_by_severity"]):
                res_rows.append([_para(sev, body_style), _para(str(retest_sum["residual_by_severity"][sev]), body_style)])
            res_table = Table(res_rows, colWidths=[150, 80], repeatRows=1)
            res_table.setStyle(summary_table_style)
            story.append(res_table)
    story.append(Spacer(1, 15))

    for f in findings:
        story.append(_para(f"Finding: {f.get('title', '')}", h1_style))
        story.append(Paragraph(
            f"<b>Severity:</b> {_pdf_text(f.get('severity', ''))} | "
            f"<b>Status:</b> {_pdf_text(f.get('status', 'Draft'))} | "
            f"<b>Category:</b> {_pdf_text(f.get('category', 'N/A'))}", body_style))
        story.append(Paragraph(
            f"<b>Affected:</b> {_pdf_text(f.get('affected_url') or f.get('affected_host') or 'N/A')} | "
            f"<b>Environment:</b> {_pdf_text(f.get('environment', 'N/A'))}", body_style))
        story.append(Paragraph(
            f"<b>HTTP Method:</b> {_pdf_text(f.get('http_method', 'N/A'))} | "
            f"<b>Parameter:</b> {_pdf_text(f.get('parameter', 'N/A'))}", body_style))
        story.append(Paragraph(
            f"<b>CWE:</b> {_pdf_text(f.get('cwe', ''))} | "
            f"<b>CVSS:</b> {_pdf_text(f.get('cvss', ''))}", body_style))

        flags = qa_flag_text(f)
        if flags:
            story.append(Paragraph(f"<b>VERIFICATION FLAGS:</b> {_pdf_text(flags)}", qa_flag_style))

        # Risk priority + framework mapping
        assessment = risk_map.assess(f)
        risk = assessment["risk"]
        fw = assessment["frameworks"]
        story.append(Paragraph("<b>Risk Priority &amp; Framework Mapping:</b>", h2_style))
        hexc = _RISK_HEX.get(risk["priority"], "#2d3748")
        story.append(Paragraph(
            f"<b>Risk priority:</b> <font color='{hexc}'><b>{risk['priority']}</b></font> ({risk['score']}/100)",
            body_style))
        if risk["rationale"]:
            story.append(_para(f"Rationale: {risk['rationale']}", body_style))
        if fw.get("mapped"):
            story.append(Paragraph(f"<b>OWASP Top 10:2025:</b> {_pdf_text(fw['owasp'])}", body_style))
            story.append(Paragraph(f"<b>PCI DSS:</b> {_pdf_text(fw['pci'])}", body_style))
            story.append(Paragraph(f"<b>NIST SP 800-53:</b> {_pdf_text(fw['nist'])}", body_style))
            if fw.get("attack"):
                story.append(Paragraph(f"<b>MITRE ATT&amp;CK:</b> {_pdf_text(fw['attack'])}", body_style))
            story.append(_para("Framework mappings are indicative; confirm against the engagement scope.", body_style))
        else:
            story.append(_para(f"Framework mapping: unmapped ({fw.get('basis', '')}) - assign manually.", body_style))

        story.append(Paragraph("<b>Description:</b>", h2_style))
        story.append(_para(f.get("description", ""), body_style))

        story.append(Paragraph("<b>Evidence:</b>", h2_style))
        story.append(_para(f.get("evidence", ""), evidence_style))

        if f.get("evidence_files"):
            story.append(Paragraph(f"<b>Evidence Files:</b> {_pdf_text(f.get('evidence_files'))}", body_style))

        story.append(Paragraph("<b>Impact:</b>", h2_style))
        story.append(_para(f.get("impact", ""), body_style))

        story.append(Paragraph("<b>Exploitation Scenario:</b>", h2_style))
        story.append(_para(f.get("scenario", ""), body_style))

        story.append(Paragraph("<b>Reproduction Steps:</b>", h2_style))
        story.append(_para(f.get("steps", ""), body_style))

        story.append(Paragraph("<b>Remediation Roadmap:</b>", h2_style))
        story.append(_para(f.get("remediation", ""), body_style))

        # Retest detail (only when retested)
        if _retested(f):
            story.append(Paragraph("<b>Retest:</b>", h2_style))
            story.append(Paragraph(
                f"<b>Status:</b> {_pdf_text(f.get('retest_status') or 'N/A')} | "
                f"<b>Round:</b> {_pdf_text(f.get('retest_round') or 0)} | "
                f"<b>Date:</b> {_pdf_text(f.get('retest_date') or 'N/A')} | "
                f"<b>Retester:</b> {_pdf_text(f.get('retester') or 'N/A')}", body_style))
            story.append(Paragraph(
                f"<b>Original severity:</b> {_pdf_text(f.get('original_severity') or f.get('severity') or 'N/A')} | "
                f"<b>Current severity:</b> {_pdf_text(f.get('severity') or 'N/A')} | "
                f"<b>First found:</b> {_pdf_text(f.get('first_found_date') or f.get('created_at') or 'N/A')}", body_style))
            if f.get("retest_evidence"):
                story.append(Paragraph("<b>Retest evidence:</b>", h2_style))
                story.append(_para(f.get("retest_evidence", ""), evidence_style))
            history = _retest_history(f)
            if history:
                story.append(Paragraph("<b>Retest history:</b>", body_style))
                for ev in history:
                    line = f"R{ev.get('round')} - {ev.get('date')} - {ev.get('status')}"
                    if ev.get("retester"):
                        line += f" - {ev.get('retester')}"
                    if ev.get("note"):
                        line += f" - {ev.get('note')}"
                    story.append(_para(line, body_style))

        story.append(Paragraph("<b>False Positive Checks:</b>", h2_style))
        story.append(_para(f.get("fp_checks", ""), body_style))

        if f.get("retest_notes"):
            story.append(Paragraph("<b>Retest Notes:</b>", h2_style))
            story.append(_para(f.get("retest_notes", ""), body_style))

        refs = f.get("references_data") or f.get("references") or ""
        if refs:
            story.append(Paragraph("<b>References:</b>", h2_style))
            story.append(_para(refs, body_style))

        qa_text = (f.get("additional_remarks") or "").strip()
        if qa_text:
            story.append(Paragraph("<b>Automated QA / Verification:</b>", h2_style))
            for line in qa_text.splitlines():
                if not line.strip():
                    continue
                story.append(_para(line, qa_warn_style if is_warning_line(line) else body_style))

        story.append(Spacer(1, 12))

    doc.build(story)


# ===========================================================================
# XLSX
# ===========================================================================
def _xlsx_row(ws, values, bold=False, fill=None):
    ws.append(values)
    if bold or fill:
        r = ws.max_row
        for c in range(1, len(values) + 1):
            cell = ws.cell(row=r, column=c)
            if bold:
                cell.font = Font(bold=True)
            if fill:
                cell.fill = fill


def export_to_excel(project: dict, findings: list, filepath: str):
    wb = Workbook()
    ws = wb.active
    ws.title = "Vulnerability Summary"

    headers = [
        "ID", "Title", "Severity", "Risk Priority", "Risk Score", "Status", "Category", "Environment",
        "Affected Host", "Affected URL", "HTTP Method", "Parameter", "Owner",
        "CWE", "CVSS", "OWASP 2025", "PCI DSS", "NIST 800-53", "MITRE ATT&CK",
        "Description", "Evidence", "Evidence Files", "Impact", "Remediation",
        "False Positive Checks",
        "Retest Status", "Retest Round", "Retest Date", "Retester", "Retest Evidence", "Retest Notes",
        "Risk Rationale", "QA Flags", "Verification Notes",
    ]
    ws.append(headers)

    header_font = Font(name="Arial", size=11, bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="1A365D", end_color="1A365D", fill_type="solid")
    thin_border = Border(left=Side(style="thin", color="CCCCCC"), right=Side(style="thin", color="CCCCCC"),
                         top=Side(style="thin", color="CCCCCC"), bottom=Side(style="thin", color="CCCCCC"))
    danger_fill = PatternFill(start_color="F8D7DA", end_color="F8D7DA", fill_type="solid")
    warn_fill = PatternFill(start_color="FFF3CD", end_color="FFF3CD", fill_type="solid")
    qa_flags_col = headers.index("QA Flags") + 1
    priority_col = headers.index("Risk Priority") + 1

    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")

    for i, f in enumerate(findings, 2):
        qa = summarize_qa(f)
        assessment = risk_map.assess(f)
        risk = assessment["risk"]
        fw = assessment["frameworks"]
        ws.append([
            i - 1, f.get("title"), f.get("severity"), risk["priority"], risk["score"],
            f.get("status"), f.get("category"), f.get("environment"),
            f.get("affected_host"), f.get("affected_url"), f.get("http_method"), f.get("parameter"), f.get("owner"),
            f.get("cwe"), f.get("cvss"),
            fw.get("owasp", ""), fw.get("pci", ""), fw.get("nist", ""), fw.get("attack", ""),
            f.get("description"), f.get("evidence"), f.get("evidence_files"), f.get("impact"), f.get("remediation"),
            f.get("fp_checks"),
            f.get("retest_status") or "Not Retested", f.get("retest_round") or 0, f.get("retest_date"),
            f.get("retester"), f.get("retest_evidence"), f.get("retest_notes"),
            risk["rationale"], qa_flag_text(f), f.get("additional_remarks"),
        ])
        for col_num in range(1, len(headers) + 1):
            c = ws.cell(row=i, column=col_num)
            c.border = thin_border
            c.alignment = Alignment(vertical="top", wrap_text=True)
        if qa["level"] == "danger":
            ws.cell(row=i, column=qa_flags_col).fill = danger_fill
        elif qa["level"] == "warn":
            ws.cell(row=i, column=qa_flags_col).fill = warn_fill
        prio_fill = _RISK_XLSX_FILL.get(risk["priority"])
        if prio_fill:
            ws.cell(row=i, column=priority_col).fill = prio_fill

    # Summary sheet: risk priorities, OWASP coverage, remediation status
    risk_sum = _risk_summary(findings)
    retest_sum = _retest_summary(findings)
    ws2 = wb.create_sheet("Summary")

    _xlsx_row(ws2, ["Risk Priority", "Findings"], bold=True, fill=header_fill)
    for pr in _PRIORITY_ORDER:
        if risk_sum["priorities"].get(pr):
            _xlsx_row(ws2, [pr, risk_sum["priorities"][pr]], fill=_RISK_XLSX_FILL.get(pr))

    _xlsx_row(ws2, [])
    _xlsx_row(ws2, ["OWASP Top 10:2025 Category", "Findings"], bold=True, fill=header_fill)
    for cat in sorted(risk_sum["owasp"]):
        _xlsx_row(ws2, [cat, risk_sum["owasp"][cat]])
    if risk_sum["unmapped"]:
        _xlsx_row(ws2, ["Unmapped (assign manually)", risk_sum["unmapped"]])

    if retest_sum["any_retested"]:
        _xlsx_row(ws2, [])
        _xlsx_row(ws2, ["Remediation Status", "Value"], bold=True, fill=header_fill)
        _xlsx_row(ws2, ["Total findings", retest_sum["total"]])
        _xlsx_row(ws2, ["Retested", retest_sum["retested"]])
        _xlsx_row(ws2, ["Fixed", retest_sum["fixed"]])
        _xlsx_row(ws2, ["Still open", retest_sum["open"]])
        _xlsx_row(ws2, ["Accepted risk", retest_sum["accepted"]])
        _xlsx_row(ws2, ["Not retested", retest_sum["not_retested"]])
        _xlsx_row(ws2, ["Remediation rate (%)", round(retest_sum["remediation_rate"], 1)])
        if retest_sum["residual_by_severity"]:
            _xlsx_row(ws2, [])
            _xlsx_row(ws2, ["Residual Risk - Open by Original Severity", "Open"], bold=True, fill=header_fill)
            for sev in _order_severity(retest_sum["residual_by_severity"]):
                _xlsx_row(ws2, [sev, retest_sum["residual_by_severity"][sev]])

    wb.save(filepath)


# ===========================================================================
# JSON
# ===========================================================================
def export_to_json(project: dict, findings: list, exec_summary: str, methodology: str, filepath: str):
    enriched = []
    for f in findings:
        item = dict(f)
        item["risk_priority"] = risk_map.compute_risk_priority(f)
        item["framework_mapping"] = risk_map.map_frameworks(f)
        enriched.append(item)

    output = {
        "project": project,
        "executive_summary": exec_summary,
        "methodology": methodology,
        "risk_summary": _risk_summary(findings),
        "retest_summary": _retest_summary(findings),
        "findings": enriched,
    }
    with open(filepath, "w", encoding="utf-8") as handle:
        json.dump(output, handle, indent=4)
